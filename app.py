import streamlit as st
import google.generativeai as genai
import os, io, requests, time
import pandas as pd
from PIL import Image
from PyPDF2 import PdfReader
from docx import Document
from bs4 import BeautifulSoup
from urllib.parse import urljoin

# --- CẤU HÌNH ---
st.set_page_config(page_title="Hanzi Intelligence Pro v4", page_icon="🐲", layout="wide")

st.markdown("""
    <style>
    .main { background-color: #fdfaf6; }
    .stButton>button { background: linear-gradient(45deg, #2c3e50, #c0392b); color: white; border-radius: 10px; font-weight: bold; }
    .stDataFrame { background-color: white; border-radius: 10px; }
    .lesson-box { padding: 20px; border-radius: 10px; border-left: 10px solid #c0392b; background-color: #ffffff; box-shadow: 2px 2px 10px rgba(0,0,0,0.1); margin-bottom: 20px; }
    </style>
    """, unsafe_allow_html=True)

# --- KẾT NỐI API ---
try:
    api_key = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=api_key)
    available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods and 'gemini' in m.name]
except:
    st.error("⚠️ Vui lòng cấu hình GEMINI_API_KEY.")
    st.stop()

# --- HÀM HỖ TRỢ ---
def get_html(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        res = requests.get(url, headers=headers, timeout=15)
        res.encoding = res.apparent_encoding
        return res.text
    except: return None

def save_docx(content, title):
    doc = Document()
    doc.add_heading(title, 0)
    for line in content.split('\n'):
        if line.strip(): doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio

# --- SIDEBAR ---
with st.sidebar:
    st.title("🐲 SIÊU NÃO BỘ V4")
    selected_model = st.selectbox("🎯 Chọn Bộ Não:", available_models, index=0)
    st.divider()
    menu = st.radio("🚀 CHỌN CHẾ ĐỘ:", [
        "🔍 Thợ Săn Truyện & Lọc Top",
        "🏭 Dịch Hàng Loạt Theo Bộ",
        "🎓 Giáo Trình Tự Động",
        "🧠 Đại Sư Phân Tích (Upload)",
        "🖼️ Dịch Ảnh OCR"
    ])
    st.divider()
    st.info("Phiên bản v4: Tích hợp Scraper thông minh và Lọc truyện.")

model = genai.GenerativeModel(selected_model)

# --- 1. THỢ SĂN TRUYỆN & LỌC TOP ---
if menu == "🔍 Thợ Săn Truyện & Lọc Top":
    st.title("🔍 Thợ Săn Truyện: Quét, Phân Loại & Lọc Top")
    st.write("Dán link trang danh mục hoặc bảng xếp hạng của web truyện (Vd: qidian.com, 69shuba...).")
    
    catalog_url = st.text_input("Link trang danh mục/Xếp hạng:")
    
    if st.button("🚀 Quét Toàn Bộ Danh Mục"):
        with st.spinner("AI đang 'đọc' website và nhặt dữ liệu..."):
            html = get_html(catalog_url)
            if html:
                prompt = f"""
                Từ mã nguồn HTML này, hãy tìm và trích xuất danh sách tất cả các bộ truyện.
                Đối với mỗi bộ truyện, hãy tìm các thông tin sau:
                1. Tên truyện (Dịch sang Tiếng Việt).
                2. Đường link dẫn đến trang giới thiệu bộ truyện.
                3. Thể loại (Tiên hiệp, Đô thị, vv).
                4. Số chương hiện có.
                5. Lượt xem (Views) và Đánh giá (Rating/Score).
                
                Hãy trình bày kết quả dưới dạng một Bảng dữ liệu Markdown rõ ràng. 
                Sắp xếp theo thứ tự Lượt xem hoặc Đánh giá từ cao xuống thấp.
                
                MÃ NGUỒN HTML:
                {html[:25000]}
                """
                res = model.generate_content(prompt)
                st.markdown(res.text)
                st.success("Mẹo: Bạn hãy copy Link chương 1 của bộ truyện muốn dịch để sang bước tiếp theo.")
            else:
                st.error("Không thể kết nối đến website.")

# --- 2. DỊCH HÀNG LOẠT THEO BỘ ---
elif menu == "🏭 Dịch Hàng Loạt Theo Bộ":
    st.title("🏭 Cỗ Máy Dịch Thuật Công Nghiệp")
    st.info("AI sẽ tự động tìm nút 'Chương sau' để dịch liên tiếp mà không cần dán từng link.")
    
    col1, col2 = st.columns(2)
    with col1:
        start_url = st.text_input("Link chương bắt đầu (Chương 1):")
        num_chaps = st.number_input("Số chương muốn dịch:", 1, 500, 10)
    with col2:
        style_req = st.text_area("Yêu cầu văn phong (Vd: Dịch thoát ý, giữ Hán Việt):", "Dịch tiên hiệp, xưng hô Ta - Ngươi, giữ nguyên tên riêng.")
        glossary = st.text_area("Từ điển (Mỗi dòng 1 từ):", "Trúc Cơ, Linh Khí")

    if st.button("🚀 BẮT ĐẦU CHIẾN DỊCH DỊCH THUẬT"):
        full_content = ""
        current_url = start_url
        p_bar = st.progress(0)
        
        for i in range(num_chaps):
            st.write(f"📂 Đang cào chương {i+1}: {current_url}")
            html = get_html(current_url)
            if not html: break
            
            # AI thực hiện 3 việc: Lấy nội dung, Tìm link sau, và Dịch luôn
            prompt = f"""
            Nhiệm vụ:
            1. Trích nội dung chương truyện (bỏ qua quảng cáo).
            2. Tìm URL của chương tiếp theo.
            3. Dịch nội dung sang tiếng Việt: {style_req}. Thuật ngữ: {glossary}.
            
            Định dạng trả về:
            CONTENT: [Bản dịch]
            NEXT_URL: [Link sau]
            
            HTML: {html[:20000]}
            """
            try:
                res_raw = model.generate_content(prompt).text
                chapter_val = res_raw.split("CONTENT:")[1].split("NEXT_URL:")[0].strip()
                next_url = res_raw.split("NEXT_URL:")[1].strip()
                
                full_content += f"\n\n--- CHƯƠNG {i+1} ---\n\n" + chapter_val
                
                # Cập nhật cho vòng lặp sau
                if next_url.startswith("http"):
                    current_url = next_url
                else:
                    current_url = urljoin(current_url, next_url)
                
                st.success(f"✅ Xong chương {i+1}")
            except:
                st.error(f"Dừng lại vì lỗi cấu hình tại chương {i+1}")
                break
            
            p_bar.progress((i+1)/num_chaps)
            time.sleep(1) # Tránh bị chặn IP
            
        st.download_button("📥 Tải Trọn Bộ Word", save_docx(full_content, "Truyen_Full").getvalue(), "Truyen_Dich.docx")

# --- 3. GIÁO TRÌNH TỰ ĐỘNG ---
elif menu == "🎓 Giáo Trình Tự Động":
    st.title("🎓 Học Viện Hán Ngữ: Thiết Kế Giáo Trình")
    topic = st.text_input("Chủ đề bạn muốn học hôm nay:")
    if st.button("Tạo bài giảng"):
        prompt = f"Bạn là giáo sư ngôn ngữ. Dạy tôi '{topic}'. Gồm: Bài học, Từ vựng (Hán-Pinyin-Hán Việt-Nghĩa), Ngữ pháp, Cách viết chữ và Bài tập."
        res = model.generate_content(prompt)
        st.markdown("<div class='lesson-box'>", unsafe_allow_html=True)
        st.markdown(res.text)
        st.markdown("</div>", unsafe_allow_html=True)

# --- 4. ĐẠI SƯ PHÂN TÍCH ---
elif menu == "🧠 Đại Sư Phân Tích (Upload)":
    st.title("🧠 Chuyên Gia Quy Nạp & Giảng Giải")
    files = st.file_uploader("Nạp sách/tài liệu:", accept_multiple_files=True)
    q = st.text_input("Câu hỏi về nội dung sách:")
    if st.button("Phân Tích Chuyên Sâu") and files:
        # Logic đọc file tương tự bản trước
        st.write("AI đang nghiên cứu...")

# --- 5. DỊCH ẢNH OCR ---
elif menu == "🖼️ Dịch Ảnh OCR":
    st.title("📸 Dịch Sách & Truyện Qua Ảnh")
    imgs = st.file_uploader("Tải ảnh:", accept_multiple_files=True)
    if st.button("Dịch Ảnh Batch") and imgs:
        for im_f in imgs:
            img = Image.open(im_f)
            st.image(img, width=300)
            res = model.generate_content(["Đọc chữ (kể cả dọc) và dịch sang TV mượt mà:", img])
            st.markdown(res.text)
